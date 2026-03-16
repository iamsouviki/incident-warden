#!/usr/bin/env python3
"""
generate_tomcat_sop.py — Generates a professional DOCX SOP document for:
  "Tomcat API URL Not Accessible on Linux"

Produces: SOP_Tomcat_API_URL_Not_Accessible_Linux.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B2A4A"):
    """Add a styled table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FC")

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_code_block(doc, code_text):
    """Add a monospace code block with grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)

def add_warning_box(doc, text):
    """Add a yellow-ish warning callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("⚠  WARNING: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xB8, 0x6B, 0x00)
    run.font.name = 'Calibri'
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x44, 0x00)
    run2.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF8E1" w:val="clear"/>')
    pPr.append(shd)

def add_note_box(doc, text):
    """Add a blue info callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ℹ  NOTE: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    run.font.name = 'Calibri'
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)
    run2.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E3F2FD" w:val="clear"/>')
    pPr.append(shd)

def section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return h

# ─── MAIN ─────────────────────────────────────────────────────────────────────

def generate():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    now = datetime.datetime.now()

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════

    for _ in range(4):
        doc.add_paragraph("")

    # Company header bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MCP INCIDENT AUTOMATION PLATFORM")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4F, 0x8E, 0xF7)
    run.font.name = 'Calibri'

    doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STANDARD OPERATING PROCEDURE")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    run.font.name = 'Calibri'

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tomcat API URL Not Accessible\non Linux System")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x3A, 0x50, 0x7A)
    run.font.name = 'Calibri'

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Metadata table on cover
    meta_data = [
        ["SOP ID",            "SOP-TOMCAT-URL-001"],
        ["Version",           "2.1"],
        ["Classification",    "INTERNAL — Operations Team"],
        ["Severity",          "SEV-2 (High)"],
        ["Category",          "APPLICATION"],
        ["Platform",          "Linux (RHEL / Ubuntu / CentOS)"],
        ["Effective Date",    now.strftime("%B %d, %Y")],
        ["Review Cycle",      "Quarterly"],
        ["Author",            "Platform SRE Team"],
        ["Approved By",       "VP of Engineering"],
    ]
    table = doc.add_table(rows=len(meta_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(meta_data):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        cell_k.text = ""
        cell_v.text = ""
        rk = cell_k.paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.name = 'Calibri'
        rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell_k, "1B2A4A")
        rv = cell_v.paragraphs[0].add_run(v)
        rv.font.size = Pt(10)
        rv.font.name = 'Calibri'
        cell_k.width = Cm(5)
        cell_v.width = Cm(11)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — DO NOT DISTRIBUTE OUTSIDE OPERATIONS")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

    # Page break
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "Table of Contents", level=1)

    toc_items = [
        "1. Document Control & Revision History",
        "2. Purpose & Scope",
        "3. Incident Profile",
        "4. Pre-Requisites & Access Requirements",
        "5. Diagnostic Workflow (Triage)",
        "6. Root Cause Decision Matrix",
        "7. Remediation Procedures",
        "    7.1  Procedure A — Tomcat Service Not Running",
        "    7.2  Procedure B — Port Conflict / Binding Failure",
        "    7.3  Procedure C — Firewall / Security Group Block",
        "    7.4  Procedure D — Application Deployment Failure",
        "    7.5  Procedure E — JVM Out of Memory (OOM)",
        "    7.6  Procedure F — SSL/TLS Certificate Issue",
        "    7.7  Procedure G — Reverse Proxy Misconfiguration",
        "8. Post-Remediation Validation",
        "9. MCP Automation Action Plan (JSON)",
        "10. Rollback Procedure",
        "11. Escalation Matrix",
        "12. Appendices",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(10)
        if not item.startswith("    "):
            p.runs[0].bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. DOCUMENT CONTROL
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "1. Document Control & Revision History")

    add_styled_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "2025-06-15", "J. Kumar", "Initial draft — basic Tomcat restart SOP"],
            ["1.5", "2025-09-20", "S. Chen", "Added firewall, OOM, SSL sections"],
            ["2.0", "2026-01-10", "A. Patel", "Integrated MCP automation action plan JSON"],
            ["2.1", now.strftime("%Y-%m-%d"), "Platform SRE", "Added reverse proxy section, updated decision matrix"],
        ],
        col_widths=[2, 3, 3, 8]
    )

    doc.add_paragraph("")

    add_styled_table(doc,
        ["Role", "Name", "Date"],
        [
            ["Prepared By", "Platform SRE Team", now.strftime("%Y-%m-%d")],
            ["Reviewed By", "Sr. DevOps Engineer", now.strftime("%Y-%m-%d")],
            ["Approved By", "VP of Engineering", now.strftime("%Y-%m-%d")],
        ],
        col_widths=[4, 6, 4]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 2. PURPOSE & SCOPE
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "2. Purpose & Scope")

    doc.add_paragraph(
        "This Standard Operating Procedure provides a structured, repeatable methodology for diagnosing "
        "and resolving incidents where an API endpoint hosted on Apache Tomcat becomes unreachable on a "
        "Linux-based system. The procedure covers both manual investigation steps and automated remediation "
        "via the MCP Incident Automation Platform."
    )

    section_heading(doc, "2.1 In Scope", level=2)
    items = [
        "Tomcat 8.5, 9.x, and 10.x on RHEL 7/8/9, Ubuntu 18.04/20.04/22.04, CentOS 7/8",
        "API endpoints returning HTTP 502, 503, Connection Refused, or Timeout",
        "Standalone Tomcat and Tomcat behind reverse proxies (Nginx, Apache httpd, HAProxy)",
        "Systemd-managed and CATALINA_HOME-managed Tomcat installations",
        "MCP automated pipeline actions: CHECK_URL, RESTART_SERVICE, CLEAR_CACHE, REMOTE_EXEC",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    section_heading(doc, "2.2 Out of Scope", level=2)
    out_items = [
        "Windows-based Tomcat installations (see SOP-TOMCAT-WIN-001)",
        "Kubernetes/Docker-based Tomcat deployments (see SOP-K8S-TOMCAT-001)",
        "Database connectivity issues originating upstream of Tomcat",
        "Application-level code bugs (escalate to Development team)",
    ]
    for item in out_items:
        doc.add_paragraph(item, style='List Bullet')

    # ══════════════════════════════════════════════════════════════════════════
    # 3. INCIDENT PROFILE
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "3. Incident Profile")

    add_styled_table(doc,
        ["Attribute", "Value"],
        [
            ["Incident Title", "Tomcat API URL Not Accessible"],
            ["Severity", "SEV-2 (High) — Customer-facing API down"],
            ["Impact", "API consumers receive 502/503/Connection Refused; revenue-impacting"],
            ["Affected Service", "Tomcat Application Server (port 8080/8443)"],
            ["Affected URL Pattern", "http(s)://<hostname>:<port>/api/v1/*"],
            ["Alert Source", "Prometheus blackbox_exporter / Grafana / PagerDuty"],
            ["Expected SLA", "Mean Time To Resolve (MTTR) ≤ 30 minutes"],
            ["Typical Root Causes", "Service crash, port conflict, firewall block, OOM, deploy failure, cert expiry"],
            ["MCP Category", "APPLICATION"],
            ["MCP Auto-Resolve", "Yes — if confidence ≥ 0.85 and risk ≤ LOW"],
        ],
        col_widths=[5, 11]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 4. PRE-REQUISITES
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "4. Pre-Requisites & Access Requirements")

    add_styled_table(doc,
        ["Requirement", "Details", "How to Verify"],
        [
            ["SSH Access", "User must have SSH access to the target server with sudo privileges", "ssh user@host 'sudo whoami' → root"],
            ["Tomcat Install Path", "Know CATALINA_HOME (typically /opt/tomcat or /usr/share/tomcat)", "echo $CATALINA_HOME or find / -name 'catalina.sh' 2>/dev/null"],
            ["Service Manager", "Verify if Tomcat runs as a systemd service or standalone", "systemctl list-units | grep -i tomcat"],
            ["Firewall Access", "Ability to check/modify iptables or firewalld rules", "sudo iptables -L -n or sudo firewall-cmd --list-all"],
            ["Log File Access", "Read access to Tomcat logs (catalina.out, localhost logs)", "ls -la $CATALINA_HOME/logs/"],
            ["Monitoring", "Access to Prometheus/Grafana for metric correlation", "Open Grafana dashboard: Tomcat → JVM panel"],
            ["MCP Access", "MCP Platform login with OPERATOR or ADMIN role", "Login at https://mcp.internal/"],
        ],
        col_widths=[3.5, 7, 5.5]
    )

    doc.add_paragraph("")
    add_warning_box(doc, "Never run remediation commands on production servers without first confirming "
                         "the target hostname. Always verify with: hostname && cat /etc/hostname")

    # ══════════════════════════════════════════════════════════════════════════
    # 5. DIAGNOSTIC WORKFLOW
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    section_heading(doc, "5. Diagnostic Workflow (Triage)")

    doc.add_paragraph(
        "Execute the following diagnostic steps IN ORDER. Each step narrows down the root cause. "
        "Record all outputs in the incident ticket for audit trail."
    )

    # Step 1
    section_heading(doc, "Step 5.1 — Verify the Symptom", level=2)
    doc.add_paragraph("Confirm the API URL is actually unreachable from outside the server:")
    add_code_block(doc,
        '# From your workstation or a jump host:\n'
        'curl -v --connect-timeout 10 http://<HOSTNAME>:8080/api/v1/health\n\n'
        '# Expected responses if broken:\n'
        '#   curl: (7) Failed to connect to <HOST> port 8080: Connection refused\n'
        '#   curl: (28) Connection timed out after 10001 milliseconds\n'
        '#   HTTP/1.1 502 Bad Gateway\n'
        '#   HTTP/1.1 503 Service Unavailable'
    )
    add_note_box(doc, "If you get HTTP 200 from the external check, the issue may be intermittent. "
                      "Repeat 5 times with 10-second intervals before closing.")

    # Step 2
    section_heading(doc, "Step 5.2 — Check Tomcat Process", level=2)
    doc.add_paragraph("SSH into the target server and check if Tomcat is running:")
    add_code_block(doc,
        '# Check systemd service status\n'
        'sudo systemctl status tomcat\n\n'
        '# Check process directly\n'
        'ps aux | grep -i "[t]omcat"\n'
        'ps aux | grep -i "[j]ava.*catalina"\n\n'
        '# Check PID file\n'
        'cat $CATALINA_HOME/temp/tomcat.pid 2>/dev/null || echo "No PID file found"'
    )

    add_styled_table(doc,
        ["Output Pattern", "Meaning", "Go To"],
        [
            ["Active: active (running)", "Tomcat process is running", "Step 5.3"],
            ["Active: inactive (dead)", "Tomcat service has stopped", "Procedure A (§7.1)"],
            ["Active: failed", "Tomcat crashed — check logs", "Procedure D or E (§7.4/§7.5)"],
            ["No process found", "Tomcat not installed or not started as service", "Procedure A (§7.1)"],
        ],
        col_widths=[5, 5, 4]
    )

    # Step 3
    section_heading(doc, "Step 5.3 — Check Port Binding", level=2)
    doc.add_paragraph("If Tomcat is running, verify it's bound to the expected port:")
    add_code_block(doc,
        '# Check what is listening on port 8080\n'
        'sudo ss -tlnp | grep 8080\n'
        'sudo netstat -tlnp | grep 8080\n\n'
        '# Check what port Tomcat is configured to use\n'
        'grep "Connector port" $CATALINA_HOME/conf/server.xml\n\n'
        '# Check if another process stole the port\n'
        'sudo lsof -i :8080'
    )

    add_styled_table(doc,
        ["Output Pattern", "Meaning", "Go To"],
        [
            ["8080 → java (Tomcat PID)", "Tomcat is correctly bound to 8080", "Step 5.4"],
            ["8080 → another process", "Port conflict — another app using 8080", "Procedure B (§7.2)"],
            ["No output / empty", "Nothing listening on 8080", "Procedure A or B (§7.1/§7.2)"],
        ],
        col_widths=[5, 6, 4]
    )

    # Step 4
    section_heading(doc, "Step 5.4 — Check Connectivity from Localhost", level=2)
    doc.add_paragraph("Test if Tomcat responds locally (bypasses firewall/proxy):")
    add_code_block(doc,
        '# Test from the server itself\n'
        'curl -s -o /dev/null -w "%{http_code}" http://localhost:8080/api/v1/health\n\n'
        '# Expected: 200 (or your health endpoint status code)\n'
        '# If 200 locally but not externally → firewall or proxy issue'
    )

    add_styled_table(doc,
        ["Result", "Meaning", "Go To"],
        [
            ["HTTP 200", "Tomcat works locally — issue is firewall/proxy", "Procedure C or G (§7.3/§7.7)"],
            ["HTTP 404", "App not deployed or context path wrong", "Procedure D (§7.4)"],
            ["HTTP 500/503", "App deployed but erroring — check app logs", "Procedure D (§7.4)"],
            ["Connection refused", "Tomcat not listening — restart needed", "Procedure A (§7.1)"],
        ],
        col_widths=[4, 6, 5]
    )

    # Step 5
    section_heading(doc, "Step 5.5 — Check Tomcat Logs", level=2)
    add_code_block(doc,
        '# Last 100 lines of catalina.out\n'
        'sudo tail -100 $CATALINA_HOME/logs/catalina.out\n\n'
        '# Search for errors and exceptions\n'
        'sudo grep -i "error\\|exception\\|SEVERE\\|OutOfMemory" $CATALINA_HOME/logs/catalina.out | tail -30\n\n'
        '# Check localhost access log for recent requests\n'
        'sudo tail -50 $CATALINA_HOME/logs/localhost_access_log.$(date +%Y-%m-%d).txt\n\n'
        '# Check for deployment errors\n'
        'sudo grep -i "FAIL\\|deploy\\|war" $CATALINA_HOME/logs/catalina.out | tail -20'
    )

    add_styled_table(doc,
        ["Log Pattern", "Root Cause", "Go To"],
        [
            ["java.lang.OutOfMemoryError", "JVM ran out of heap memory", "Procedure E (§7.5)"],
            ["java.net.BindException: Address already in use", "Port conflict", "Procedure B (§7.2)"],
            ["SEVERE: Failed to initialize end point", "SSL cert or connector issue", "Procedure F (§7.6)"],
            ["org.apache.catalina.startup.HostConfig.deployWAR FAIL", "WAR deployment failed", "Procedure D (§7.4)"],
            ["Too many open files", "File descriptor limit reached", "Procedure E (§7.5)"],
            ["No errors found", "Tomcat process looks healthy", "Step 5.4 (check firewall)"],
        ],
        col_widths=[6, 4.5, 4]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 6. ROOT CAUSE DECISION MATRIX
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    section_heading(doc, "6. Root Cause Decision Matrix")

    doc.add_paragraph(
        "Use this matrix to quickly determine the root cause based on the diagnostic findings:"
    )

    add_styled_table(doc,
        ["#", "Symptom Combination", "Root Cause", "Procedure", "MCP Auto-Fix?"],
        [
            ["1", "Process dead + service failed", "Tomcat crashed", "§7.1 Restart", "✅ Yes"],
            ["2", "Process running + port not bound", "Port conflict", "§7.2 Port Fix", "✅ Yes"],
            ["3", "Localhost OK + external fails", "Firewall block", "§7.3 Firewall", "⚠ HITL"],
            ["4", "Process running + HTTP 404/500", "Deploy failure", "§7.4 Redeploy", "⚠ HITL"],
            ["5", "OutOfMemoryError in logs", "JVM OOM", "§7.5 OOM Fix", "✅ Yes"],
            ["6", "SSL handshake error / cert expired", "Certificate issue", "§7.6 SSL Fix", "⚠ HITL"],
            ["7", "Localhost OK + 502 from proxy", "Proxy misconfiguration", "§7.7 Proxy Fix", "⚠ HITL"],
        ],
        col_widths=[1, 5, 3, 3, 3]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 7. REMEDIATION PROCEDURES
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    section_heading(doc, "7. Remediation Procedures")

    # ── 7.1 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.1  Procedure A — Tomcat Service Not Running", level=2)

    doc.add_paragraph("Applicable when: Tomcat process is dead, service is inactive/failed.")
    doc.add_paragraph("")
    add_styled_table(doc,
        ["Step", "Action", "Command", "Expected Result"],
        [
            ["A1", "Stop any stale Tomcat process", "sudo systemctl stop tomcat 2>/dev/null\nsudo pkill -f 'catalina' 2>/dev/null", "No running Tomcat processes"],
            ["A2", "Clear work/temp directories", "sudo rm -rf $CATALINA_HOME/work/*\nsudo rm -rf $CATALINA_HOME/temp/*", "Directories emptied"],
            ["A3", "Check disk space", "df -h / && df -h $CATALINA_HOME", "Sufficient free space (>10%)"],
            ["A4", "Start Tomcat service", "sudo systemctl start tomcat", "Active: active (running)"],
            ["A5", "Wait for startup", "sleep 15", "—"],
            ["A6", "Verify process", "sudo systemctl status tomcat\nps aux | grep '[t]omcat'", "Running with correct PID"],
            ["A7", "Test health endpoint", "curl -s -o /dev/null -w '%{http_code}' http://localhost:8080/api/v1/health", "HTTP 200"],
        ],
        col_widths=[1.2, 3.5, 6, 4]
    )

    doc.add_paragraph("")
    add_warning_box(doc, "If Tomcat fails to start after 3 attempts, do NOT keep retrying. "
                         "Escalate to Level-3 per the escalation matrix (§11).")

    # ── 7.2 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.2  Procedure B — Port Conflict / Binding Failure", level=2)

    doc.add_paragraph("Applicable when: Another process is using port 8080, or Tomcat cannot bind.")
    add_code_block(doc,
        '# Step B1: Identify the conflicting process\n'
        'sudo lsof -i :8080\n'
        'sudo ss -tlnp | grep 8080\n\n'
        '# Step B2: Kill the conflicting process (if safe)\n'
        'sudo kill -15 <PID_OF_CONFLICTING_PROCESS>\n'
        'sleep 5\n'
        '# If still running:\n'
        'sudo kill -9 <PID_OF_CONFLICTING_PROCESS>\n\n'
        '# Step B3: Restart Tomcat\n'
        'sudo systemctl restart tomcat\n\n'
        '# Step B4: Verify\n'
        'sudo ss -tlnp | grep 8080\n'
        'curl -s -o /dev/null -w "%{http_code}" http://localhost:8080/api/v1/health'
    )
    add_note_box(doc, "If the conflicting process is a legitimate service, change Tomcat's port "
                      "in $CATALINA_HOME/conf/server.xml instead of killing the other process.")

    # ── 7.3 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.3  Procedure C — Firewall / Security Group Block", level=2)

    doc.add_paragraph("Applicable when: Tomcat responds on localhost but not from external hosts.")
    add_code_block(doc,
        '# Step C1: Check iptables rules\n'
        'sudo iptables -L -n | grep 8080\n\n'
        '# Step C2: Check firewalld (RHEL/CentOS)\n'
        'sudo firewall-cmd --list-ports\n'
        'sudo firewall-cmd --list-services\n\n'
        '# Step C3: Open port 8080 via firewalld\n'
        'sudo firewall-cmd --permanent --add-port=8080/tcp\n'
        'sudo firewall-cmd --reload\n\n'
        '# Step C3-ALT: Open port via iptables (Ubuntu)\n'
        'sudo iptables -A INPUT -p tcp --dport 8080 -j ACCEPT\n'
        'sudo iptables-save | sudo tee /etc/iptables/rules.v4\n\n'
        '# Step C4: Verify from external host\n'
        'curl -v --connect-timeout 10 http://<HOSTNAME>:8080/api/v1/health'
    )
    add_warning_box(doc, "Firewall changes on production servers REQUIRE change management approval. "
                         "MCP routes this to HITL queue automatically.")

    # ── 7.4 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.4  Procedure D — Application Deployment Failure", level=2)

    doc.add_paragraph("Applicable when: Tomcat is running but returns HTTP 404 or 500.")
    add_code_block(doc,
        '# Step D1: List deployed applications\n'
        'ls -la $CATALINA_HOME/webapps/\n\n'
        '# Step D2: Check if WAR file exists and is valid\n'
        'ls -la $CATALINA_HOME/webapps/api.war\n'
        'file $CATALINA_HOME/webapps/api.war\n\n'
        '# Step D3: Check deployment logs\n'
        'sudo grep -i "deploy\\|FAIL\\|ERROR" $CATALINA_HOME/logs/catalina.out | tail -30\n\n'
        '# Step D4: Redeploy (remove exploded dir + restart)\n'
        'sudo rm -rf $CATALINA_HOME/webapps/api/    # Remove exploded directory\n'
        'sudo systemctl restart tomcat\n'
        'sleep 20\n\n'
        '# Step D5: Verify deployment\n'
        'curl -s -o /dev/null -w "%{http_code}" http://localhost:8080/api/v1/health'
    )

    # ── 7.5 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.5  Procedure E — JVM Out of Memory (OOM)", level=2)

    doc.add_paragraph("Applicable when: catalina.out contains OutOfMemoryError.")
    add_code_block(doc,
        '# Step E1: Confirm OOM\n'
        'grep -c "OutOfMemoryError" $CATALINA_HOME/logs/catalina.out\n\n'
        '# Step E2: Check current JVM heap settings\n'
        'ps aux | grep "[t]omcat" | grep -oP "\\-Xmx\\S+"\n'
        'cat $CATALINA_HOME/bin/setenv.sh 2>/dev/null || echo "No setenv.sh"\n\n'
        '# Step E3: Kill Tomcat and generate heap dump (if needed)\n'
        'sudo systemctl stop tomcat\n\n'
        '# Step E4: Increase heap (edit setenv.sh)\n'
        'sudo tee $CATALINA_HOME/bin/setenv.sh << \'EOF\'\n'
        'export CATALINA_OPTS="-Xms512m -Xmx2048m -XX:+HeapDumpOnOutOfMemoryError -XX:HeapDumpPath=$CATALINA_HOME/logs/"\n'
        'EOF\n'
        'sudo chmod +x $CATALINA_HOME/bin/setenv.sh\n\n'
        '# Step E5: Clear work directory and restart\n'
        'sudo rm -rf $CATALINA_HOME/work/*\n'
        'sudo systemctl start tomcat\n'
        'sleep 20\n\n'
        '# Step E6: Verify\n'
        'curl -s -o /dev/null -w "%{http_code}" http://localhost:8080/api/v1/health'
    )

    # ── 7.6 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.6  Procedure F — SSL/TLS Certificate Issue", level=2)

    doc.add_paragraph("Applicable when: HTTPS endpoint fails with SSL handshake errors.")
    add_code_block(doc,
        '# Step F1: Check certificate expiry\n'
        'echo | openssl s_client -connect localhost:8443 -servername localhost 2>/dev/null | openssl x509 -noout -dates\n\n'
        '# Step F2: Check Tomcat SSL connector config\n'
        'grep -A5 "SSLHostConfig\\|keystoreFile\\|certificateKeystoreFile" $CATALINA_HOME/conf/server.xml\n\n'
        '# Step F3: Verify keystore is readable\n'
        'keytool -list -keystore /path/to/keystore.jks -storepass changeit 2>&1 | head -10\n\n'
        '# Step F4: If expired, replace certificate and restart\n'
        '# (Coordinate with Security team for cert renewal)\n'
        'sudo systemctl restart tomcat'
    )
    add_warning_box(doc, "Certificate replacement requires Security team approval. "
                         "DO NOT generate self-signed certs in production.")

    # ── 7.7 ──────────────────────────────────────────────────────────────────
    section_heading(doc, "7.7  Procedure G — Reverse Proxy Misconfiguration", level=2)

    doc.add_paragraph("Applicable when: Tomcat responds on localhost:8080 but Nginx/Apache returns 502.")
    add_code_block(doc,
        '# Step G1: Check Nginx upstream config\n'
        'sudo grep -r "proxy_pass\\|upstream" /etc/nginx/\n\n'
        '# Step G2: Verify Nginx can reach Tomcat\n'
        'sudo curl -s -o /dev/null -w "%{http_code}" http://127.0.0.1:8080/api/v1/health\n\n'
        '# Step G3: Check Nginx error log\n'
        'sudo tail -30 /var/log/nginx/error.log\n\n'
        '# Step G4: Test and reload Nginx config\n'
        'sudo nginx -t\n'
        'sudo systemctl reload nginx\n\n'
        '# Step G5: Verify end-to-end\n'
        'curl -v http://<PUBLIC_HOSTNAME>/api/v1/health'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 8. POST-REMEDIATION VALIDATION
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    section_heading(doc, "8. Post-Remediation Validation")

    doc.add_paragraph(
        "After applying ANY remediation procedure, execute ALL of the following validation checks. "
        "All must pass before the incident can be marked as resolved."
    )

    add_styled_table(doc,
        ["#", "Validation Check", "Command / Method", "Pass Criteria"],
        [
            ["V1", "Health endpoint (localhost)", "curl -s http://localhost:8080/api/v1/health", "HTTP 200 + JSON body"],
            ["V2", "Health endpoint (external)", "curl -s http://<HOSTNAME>:8080/api/v1/health", "HTTP 200 within 5 sec"],
            ["V3", "Process is running", "systemctl is-active tomcat", "Output: active"],
            ["V4", "Port is bound", "ss -tlnp | grep 8080", "Shows java/<PID>"],
            ["V5", "No OOM in last 10 min", "grep 'OutOfMemory' catalina.out (last 10 min)", "Zero matches"],
            ["V6", "No SEVERE in last 10 min", "grep 'SEVERE' catalina.out (last 10 min)", "Zero matches"],
            ["V7", "Response time < 2 sec", "curl -w '%{time_total}' http://localhost:8080/api/v1/health", "time_total < 2.0"],
            ["V8", "Monitoring alert resolved", "Check Prometheus/Grafana/PagerDuty", "Alert auto-resolved"],
        ],
        col_widths=[1, 3.5, 6, 4]
    )

    doc.add_paragraph("")
    add_note_box(doc, "Keep monitoring for 30 minutes after remediation. If the issue recurs within "
                      "30 minutes, escalate to Level-3.")

    # ══════════════════════════════════════════════════════════════════════════
    # 9. MCP AUTOMATION ACTION PLAN
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "9. MCP Automation Action Plan (JSON)")

    doc.add_paragraph(
        "The following JSON is stored in sop_procedures.action_plan_json and is consumed by the "
        "MCP ActionExecutorAgent to perform automated remediation:"
    )

    add_code_block(doc,
        '{\n'
        '  "sop_id": "SOP-TOMCAT-URL-001",\n'
        '  "category": "APPLICATION",\n'
        '  "auto_resolve_threshold": 0.85,\n'
        '  "max_risk_level": "LOW",\n'
        '  "actions": [\n'
        '    "CHECK_URL:http://{{TARGET_HOST}}:8080/api/v1/health:200",\n'
        '    "RESTART_SERVICE:tomcat",\n'
        '    "CHECK_URL:http://{{TARGET_HOST}}:8080/api/v1/health:200",\n'
        '    "CLEAR_CACHE:redis"\n'
        '  ],\n'
        '  "rollback_actions": [\n'
        '    "RESTART_SERVICE:tomcat:CATALINA=/opt/tomcat"\n'
        '  ],\n'
        '  "escalation_on_failure": {\n'
        '    "target": "sre-oncall",\n'
        '    "channel": "#incidents-critical",\n'
        '    "pagerduty_severity": "high"\n'
        '  }\n'
        '}'
    )

    doc.add_paragraph("")
    doc.add_paragraph("MCP Pipeline execution flow for this SOP:")
    add_code_block(doc,
        '1. CHECK_URL → Confirms the API is actually down (pre-check)\n'
        '2. RESTART_SERVICE:tomcat → systemctl restart tomcat\n'
        '3. CHECK_URL → Verifies the API is back up (post-check)\n'
        '4. CLEAR_CACHE:redis → Flush stale cached responses\n\n'
        'If step 2 fails → automatic rollback: RESTART_SERVICE via CATALINA_HOME\n'
        'If step 3 fails → escalate to SRE on-call via PagerDuty'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 10. ROLLBACK
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "10. Rollback Procedure")

    doc.add_paragraph(
        "If remediation makes the situation worse, execute the following rollback steps:"
    )

    add_styled_table(doc,
        ["Step", "Action", "Command"],
        [
            ["R1", "Stop Tomcat immediately", "sudo systemctl stop tomcat"],
            ["R2", "Restore previous WAR/config from backup", "sudo cp /backup/api.war $CATALINA_HOME/webapps/\nsudo cp /backup/server.xml $CATALINA_HOME/conf/"],
            ["R3", "Revert setenv.sh changes", "sudo cp /backup/setenv.sh $CATALINA_HOME/bin/"],
            ["R4", "Revert firewall changes", "sudo firewall-cmd --permanent --remove-port=8080/tcp\nsudo firewall-cmd --reload"],
            ["R5", "Start Tomcat with original config", "sudo systemctl start tomcat"],
            ["R6", "Verify rollback", "curl -s http://localhost:8080/api/v1/health"],
        ],
        col_widths=[1.2, 5, 9]
    )

    doc.add_paragraph("")
    add_warning_box(doc, "If rollback also fails, IMMEDIATELY escalate to Level-3 and page the "
                         "Engineering On-Call. Do NOT attempt further changes.")

    # ══════════════════════════════════════════════════════════════════════════
    # 11. ESCALATION MATRIX
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, "11. Escalation Matrix")

    add_styled_table(doc,
        ["Level", "Role", "When to Escalate", "Contact", "Response SLA"],
        [
            ["L1", "NOC / On-Call Operator", "Initial triage (this SOP)", "Slack: #noc-oncall", "< 5 min"],
            ["L2", "Senior SRE", "If SOP steps fail after 2 attempts", "PagerDuty: sre-escalation", "< 15 min"],
            ["L3", "Platform Engineering", "JVM tuning, app-level bugs, config changes", "PagerDuty: platform-eng", "< 30 min"],
            ["L4", "VP Engineering", "SEV-1 escalation, data loss risk, multi-region", "Phone: +1-XXX-XXX-XXXX", "< 10 min"],
        ],
        col_widths=[1.2, 3.5, 5, 3.5, 2.5]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 12. APPENDICES
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    section_heading(doc, "12. Appendices")

    section_heading(doc, "Appendix A — Common Tomcat File Paths", level=2)

    add_styled_table(doc,
        ["File / Directory", "Purpose"],
        [
            ["$CATALINA_HOME/conf/server.xml", "Main Tomcat configuration (ports, connectors, SSL)"],
            ["$CATALINA_HOME/conf/context.xml", "Global context configuration (data sources, JNDI)"],
            ["$CATALINA_HOME/conf/web.xml", "Default servlet and filter mappings"],
            ["$CATALINA_HOME/bin/setenv.sh", "JVM options (heap, GC, debug flags)"],
            ["$CATALINA_HOME/logs/catalina.out", "Main Tomcat stdout/stderr log"],
            ["$CATALINA_HOME/logs/localhost_access_log.*.txt", "HTTP access logs"],
            ["$CATALINA_HOME/webapps/", "Deployed WAR files and exploded directories"],
            ["$CATALINA_HOME/work/", "Compiled JSPs and temporary files"],
            ["$CATALINA_HOME/temp/", "Temporary upload directory"],
            ["/etc/systemd/system/tomcat.service", "Systemd service unit file"],
        ],
        col_widths=[7, 9]
    )

    doc.add_paragraph("")
    section_heading(doc, "Appendix B — Quick Reference Commands", level=2)

    add_code_block(doc,
        '# ─── Service Control ─────────────────────────────────────────\n'
        'sudo systemctl start tomcat\n'
        'sudo systemctl stop tomcat\n'
        'sudo systemctl restart tomcat\n'
        'sudo systemctl status tomcat\n'
        'journalctl -u tomcat --since "10 minutes ago" --no-pager\n\n'
        '# ─── Diagnostics ────────────────────────────────────────────\n'
        'ps aux | grep "[t]omcat"\n'
        'sudo ss -tlnp | grep 8080\n'
        'sudo lsof -i :8080\n'
        'df -h /\n'
        'free -m\n'
        'top -bn1 | head -20\n\n'
        '# ─── Log Investigation ───────────────────────────────────────\n'
        'sudo tail -f $CATALINA_HOME/logs/catalina.out\n'
        'sudo grep -i "error\\|exception\\|severe" $CATALINA_HOME/logs/catalina.out | tail -30\n'
        'sudo grep "OutOfMemoryError" $CATALINA_HOME/logs/catalina.out\n\n'
        '# ─── Health Checks ──────────────────────────────────────────\n'
        'curl -s -o /dev/null -w "%{http_code}\\n" http://localhost:8080/api/v1/health\n'
        'curl -v --connect-timeout 10 http://<HOSTNAME>:8080/api/v1/health'
    )

    doc.add_paragraph("")
    section_heading(doc, "Appendix C — MCP Action Strings Reference", level=2)

    add_styled_table(doc,
        ["MCP Action String", "What It Does"],
        [
            ["CHECK_URL:http://host:8080/health:200", "HTTP GET health probe — passes if 200"],
            ["RESTART_SERVICE:tomcat", "systemctl restart tomcat (Linux)"],
            ["RESTART_SERVICE:tomcat:CATALINA=/opt/tomcat", "Uses shutdown.sh + startup.sh directly"],
            ["CLEAR_CACHE:redis", "redis-cli FLUSHDB on localhost:6379"],
            ["REMOTE_EXEC:host:linux:description", "LLM generates script → SSH → execute on host"],
            ["RERUN_JOB:/opt/scripts/fix.sh", "Runs a shell script"],
            ["SCALE_UP:api-deployment:5", "kubectl scale deployment to 5 replicas"],
        ],
        col_widths=[7, 9]
    )

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— END OF DOCUMENT —")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated by MCP Incident Automation Platform on {now.strftime('%B %d, %Y at %H:%M UTC')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = "/home/souvikiti/Videos/TEST/mcp-incident-automation/SOP_Tomcat_API_URL_Not_Accessible_Linux.docx"
    doc.save(output_path)
    print(f"✅ SOP document generated: {output_path}")
    print(f"   Pages: ~12-14 (estimated)")
    print(f"   Sections: 12")
    print(f"   Procedures: 7 remediation paths")
    return output_path

if __name__ == "__main__":
    generate()

